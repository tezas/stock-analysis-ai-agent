"""
Prompt Processor

This module handles reading and processing prompts from text files and Word documents.
"""

import re
import logging
from pathlib import Path
from typing import Dict, List, Optional, Union

import pandas as pd
from pydantic import BaseModel


class AnalysisRequest(BaseModel):
    """Analysis request model."""
    
    symbol: str
    analysis_type: str
    timeframe: Optional[str] = None
    indicators: List[str] = []
    additional_requirements: List[str] = []


class PromptProcessor:
    """
    Process prompts from various sources (text files, Word documents).
    
    This class handles reading prompts from different file formats
    and extracting structured analysis requests.
    """
    
    def __init__(self):
        """Initialize the prompt processor."""
        self.logger = logging.getLogger(__name__)
        
        # Common stock symbols for validation
        self.common_symbols = [
            'AAPL', 'MSFT', 'GOOGL', 'AMZN', 'TSLA', 'META', 'NVDA', 'NFLX',
            'AMD', 'INTC', 'CRM', 'ORCL', 'ADBE', 'PYPL', 'SQ', 'ZM',
            'SPY', 'QQQ', 'IWM', 'VTI', 'VOO', 'VEA', 'VWO'
        ]
        
        # Analysis types
        self.analysis_types = [
            'technical', 'fundamental', 'comprehensive', 'short_term',
            'long_term', 'swing_trading', 'day_trading', 'value_investing'
        ]
        
        # Technical indicators
        self.technical_indicators = [
            'sma', 'ema', 'macd', 'rsi', 'bollinger_bands', 'stochastic',
            'volume', 'moving_averages', 'support_resistance', 'trend_lines'
        ]
    
    def read_text_file(self, file_path: Union[str, Path]) -> str:
        """
        Read prompt from a text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Content of the text file
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read().strip()
            
            self.logger.info(f"Successfully read text file: {file_path}")
            return content
            
        except Exception as e:
            self.logger.error(f"Error reading text file {file_path}: {str(e)}")
            raise
    
    def read_word_document(self, file_path: Union[str, Path]) -> str:
        """
        Read prompt from a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Extracted text content
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Import docx here to avoid dependency issues
            from docx import Document
            
            doc = Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text.strip())
            
            result = '\n'.join(content).strip()
            
            self.logger.info(f"Successfully read Word document: {file_path}")
            return result
            
        except ImportError:
            self.logger.error("python-docx not installed. Please install it to read Word documents.")
            raise
        except Exception as e:
            self.logger.error(f"Error reading Word document {file_path}: {str(e)}")
            raise
    
    def process_text(self, text: str) -> str:
        """
        Process and clean text prompt.
        
        Args:
            text: Raw text prompt
            
        Returns:
            Processed text prompt
        """
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text.strip())
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\-.,!?()$%]', '', text)
        
        # Normalize common variations
        text = text.replace('stock symbol', 'symbol')
        text = text.replace('ticker', 'symbol')
        text = text.replace('analyze', 'analysis')
        
        self.logger.debug(f"Processed text: {text}")
        return text
    
    def extract_analysis_request(self, prompt: str) -> AnalysisRequest:
        """
        Extract structured analysis request from prompt.
        
        Args:
            prompt: Processed prompt text
            
        Returns:
            AnalysisRequest object
        """
        try:
            # Extract stock symbol
            symbol = self._extract_symbol(prompt)
            
            # Extract analysis type
            analysis_type = self._extract_analysis_type(prompt)
            
            # Extract timeframe
            timeframe = self._extract_timeframe(prompt)
            
            # Extract technical indicators
            indicators = self._extract_indicators(prompt)
            
            # Extract additional requirements
            additional_requirements = self._extract_additional_requirements(prompt)
            
            request = AnalysisRequest(
                symbol=symbol,
                analysis_type=analysis_type,
                timeframe=timeframe,
                indicators=indicators,
                additional_requirements=additional_requirements
            )
            
            self.logger.info(f"Extracted analysis request: {request}")
            return request
            
        except Exception as e:
            self.logger.error(f"Error extracting analysis request: {str(e)}")
            raise
    
    def _extract_symbol(self, prompt: str) -> str:
        """Extract stock symbol from prompt."""
        # Look for common patterns
        patterns = [
            r'\b([A-Z]{1,5})\s+(?:stock|symbol|ticker)',
            r'(?:analyze|analysis of|check)\s+([A-Z]{1,5})',
            r'\b([A-Z]{1,5})\s+(?:for|with|using)',
            r'\b([A-Z]{1,5})\b'  # Fallback: any 1-5 letter uppercase word
        ]
        
        for pattern in patterns:
            match = re.search(pattern, prompt, re.IGNORECASE)
            if match:
                symbol = match.group(1).upper()
                # Validate against common symbols or allow any valid symbol
                if symbol in self.common_symbols or self._is_valid_symbol(symbol):
                    return symbol
        
        # If no symbol found, try to extract from the beginning
        words = prompt.split()
        for word in words:
            if self._is_valid_symbol(word.upper()):
                return word.upper()
        
        raise ValueError("No valid stock symbol found in prompt")
    
    def _is_valid_symbol(self, symbol: str) -> bool:
        """Check if symbol is valid."""
        # Basic validation: 1-5 letters, no numbers or special chars
        return bool(re.match(r'^[A-Z]{1,5}$', symbol))
    
    def _extract_analysis_type(self, prompt: str) -> str:
        """Extract analysis type from prompt."""
        prompt_lower = prompt.lower()
        
        # Map keywords to analysis types
        type_mapping = {
            'technical': ['technical', 'chart', 'indicator', 'pattern'],
            'fundamental': ['fundamental', 'financial', 'earnings', 'revenue'],
            'comprehensive': ['comprehensive', 'full', 'complete', 'detailed'],
            'short_term': ['short term', 'short-term', 'daily', 'intraday'],
            'long_term': ['long term', 'long-term', 'monthly', 'yearly'],
            'swing_trading': ['swing', 'swing trading'],
            'day_trading': ['day trading', 'daytrading', 'intraday'],
            'value_investing': ['value', 'value investing', 'fundamental']
        }
        
        for analysis_type, keywords in type_mapping.items():
            if any(keyword in prompt_lower for keyword in keywords):
                return analysis_type
        
        # Default to comprehensive if no specific type found
        return 'comprehensive'
    
    def _extract_timeframe(self, prompt: str) -> Optional[str]:
        """Extract timeframe from prompt."""
        prompt_lower = prompt.lower()
        
        timeframe_patterns = {
            '1d': ['1 day', 'one day', 'daily', 'today'],
            '5d': ['5 days', 'five days', 'week'],
            '1mo': ['1 month', 'one month', 'monthly'],
            '3mo': ['3 months', 'three months', 'quarter'],
            '6mo': ['6 months', 'six months', 'half year'],
            '1y': ['1 year', 'one year', 'yearly', 'annual'],
            '2y': ['2 years', 'two years'],
            '5y': ['5 years', 'five years'],
            'max': ['all time', 'maximum', 'full history']
        }
        
        for timeframe, keywords in timeframe_patterns.items():
            if any(keyword in prompt_lower for keyword in keywords):
                return timeframe
        
        return None
    
    def _extract_indicators(self, prompt: str) -> List[str]:
        """Extract technical indicators from prompt."""
        prompt_lower = prompt.lower()
        indicators = []
        
        indicator_mapping = {
            'sma': ['sma', 'simple moving average', 'moving average'],
            'ema': ['ema', 'exponential moving average'],
            'macd': ['macd', 'moving average convergence divergence'],
            'rsi': ['rsi', 'relative strength index'],
            'bollinger_bands': ['bollinger', 'bollinger bands', 'bb'],
            'stochastic': ['stochastic', 'stoch'],
            'volume': ['volume', 'volume analysis'],
            'support_resistance': ['support', 'resistance', 'levels'],
            'trend_lines': ['trend', 'trend line', 'trendlines']
        }
        
        for indicator, keywords in indicator_mapping.items():
            if any(keyword in prompt_lower for keyword in keywords):
                indicators.append(indicator)
        
        return indicators
    
    def _extract_additional_requirements(self, prompt: str) -> List[str]:
        """Extract additional requirements from prompt."""
        requirements = []
        prompt_lower = prompt.lower()
        
        # Common requirements
        requirement_keywords = {
            'risk_assessment': ['risk', 'volatility', 'uncertainty'],
            'price_target': ['price target', 'target price', 'price prediction'],
            'entry_exit': ['entry', 'exit', 'buy', 'sell', 'position'],
            'comparison': ['compare', 'comparison', 'vs', 'versus'],
            'news_impact': ['news', 'earnings', 'events', 'catalyst'],
            'sector_analysis': ['sector', 'industry', 'peer', 'competitor']
        }
        
        for requirement, keywords in requirement_keywords.items():
            if any(keyword in prompt_lower for keyword in keywords):
                requirements.append(requirement)
        
        return requirements
    
    def validate_prompt(self, prompt: str) -> bool:
        """
        Validate if prompt contains required information.
        
        Args:
            prompt: Processed prompt text
            
        Returns:
            True if prompt is valid, False otherwise
        """
        try:
            # Try to extract analysis request
            request = self.extract_analysis_request(prompt)
            
            # Check if symbol is valid
            if not request.symbol:
                return False
            
            # Check if analysis type is valid
            if request.analysis_type not in self.analysis_types:
                return False
            
            return True
            
        except Exception:
            return False
    
    def enhance_prompt(self, prompt: str) -> str:
        """
        Enhance prompt with additional context and structure.
        
        Args:
            prompt: Original prompt
            
        Returns:
            Enhanced prompt
        """
        try:
            request = self.extract_analysis_request(prompt)
            
            enhanced_parts = [
                f"Analyze {request.symbol} stock",
                f"Analysis type: {request.analysis_type}"
            ]
            
            if request.timeframe:
                enhanced_parts.append(f"Timeframe: {request.timeframe}")
            
            if request.indicators:
                enhanced_parts.append(f"Technical indicators: {', '.join(request.indicators)}")
            
            if request.additional_requirements:
                enhanced_parts.append(f"Additional requirements: {', '.join(request.additional_requirements)}")
            
            enhanced_prompt = ". ".join(enhanced_parts) + "."
            
            self.logger.debug(f"Enhanced prompt: {enhanced_prompt}")
            return enhanced_prompt
            
        except Exception as e:
            self.logger.warning(f"Could not enhance prompt: {str(e)}")
            return prompt 